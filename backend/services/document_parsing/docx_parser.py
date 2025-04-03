"""
SutazAI DOCX Parser Module
Extracts text, tables, images, and identifies structural elements from DOCX documents
"""

import io
import os
import logging
import base64
import tempfile
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass

# Required dependencies (pip install python-docx lxml pillow)
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

# Configure logging
logger = logging.getLogger("docx_parser")


@dataclass
class DocxSection:
    """Data class to store information about a document section"""

    heading: str
    level: int
    content: str
    paragraphs: List[str]
    tables: List[Dict[str, Any]]


class DocxParser:
    """
    Parser for extracting information from DOCX documents including
    text, tables, images, and document structure.
    """

    def __init__(self, extract_images: bool = True):
        """
        Initialize the DOCX parser

        Args:
            extract_images: Whether to extract images from the document
        """
        self.extract_images = extract_images

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file from disk

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary with parsed information
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {"error": f"File not found: {file_path}"}

        try:
            with open(file_path, "rb") as f:
                docx_data = f.read()
            return self.parse_bytes(docx_data, filename=os.path.basename(file_path))
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            return {"error": f"Failed to parse DOCX: {str(e)}"}

    def parse_bytes(
        self, docx_data: bytes, filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse a DOCX from bytes

        Args:
            docx_data: DOCX file content as bytes
            filename: Optional filename for reference

        Returns:
            Dictionary with parsed information
        """
        try:
            # Load document from bytes
            docx_stream = io.BytesIO(docx_data)
            doc = Document(docx_stream)

            # Extract metadata
            metadata = self._extract_metadata(doc)

            # Extract content (paragraphs, tables, etc.)
            content = self._extract_content(doc)

            # Extract images if enabled
            images = []
            if self.extract_images:
                images = self._extract_images(doc, docx_data)

            # Extract document structure
            structure = self._identify_document_structure(doc, content)

            # Assemble full text
            full_text = self._assemble_full_text(content["paragraphs"])

            return {
                "success": True,
                "metadata": metadata,
                "content": content,
                "full_text": full_text,
                "structure": structure,
                "images": images,
                "filename": filename,
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX data: {str(e)}")
            return {"success": False, "error": f"Failed to parse DOCX: {str(e)}"}

    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document

        Args:
            doc: Document object

        Returns:
            Dictionary with metadata
        """
        metadata = {}

        # Extract core properties
        core_properties = doc.core_properties

        # Author information
        metadata["author"] = core_properties.author
        metadata["last_modified_by"] = core_properties.last_modified_by

        # Dates
        metadata["created"] = (
            core_properties.created.isoformat() if core_properties.created else None
        )
        metadata["modified"] = (
            core_properties.modified.isoformat() if core_properties.modified else None
        )

        # Document information
        metadata["title"] = core_properties.title
        metadata["subject"] = core_properties.subject
        metadata["keywords"] = core_properties.keywords
        metadata["category"] = core_properties.category
        metadata["comments"] = core_properties.comments
        metadata["revision"] = core_properties.revision

        # Additional metadata
        custom_properties = {}
        try:
            for prop in doc.custom_properties:  # type: ignore [attr-defined]
                custom_properties[prop.name] = prop.value
        except Exception as e:
            logger.debug(f"Could not extract custom properties: {e}")
            pass

        if custom_properties:
            metadata["custom_properties"] = custom_properties

        return metadata

    def _extract_content(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract content from DOCX document

        Args:
            doc: Document object

        Returns:
            Dictionary with content elements
        """
        content: Dict[str, List[Any]] = {
            "paragraphs": [],
            "tables": [],
            "paragraphs_with_style": [],
            "headers": [],
            "footers": [],
        }

        # Process all block items (paragraphs and tables)
        for i, block in enumerate(self._iter_block_items(doc)):
            if isinstance(block, Paragraph):
                text = block.text.strip()

                # Skip empty paragraphs
                if not text:
                    continue

                # Add to plain paragraphs
                content["paragraphs"].append(text)

                # Save paragraph with style information
                style_name = block.style.name if block.style else "Normal"
                content["paragraphs_with_style"].append(
                    {
                        "text": text,
                        "style": style_name,
                        "is_heading": style_name.startswith("Heading"),
                        "heading_level": int(style_name.replace("Heading ", ""))
                        if style_name.startswith("Heading ")
                        else 0,
                    }
                )

            elif isinstance(block, Table):
                table_data = self._extract_table(block)
                content["tables"].append(table_data)

        # Extract headers and footers if available
        try:
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            content["headers"].append(paragraph.text.strip())

                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            content["footers"].append(paragraph.text.strip())
        except Exception as e:
            logger.warning(f"Error extracting headers/footers: {str(e)}")

        return content

    def _iter_block_items(self, doc: DocxDocument):
        """
        Iterate through all paragraphs and tables in the document

        Args:
            doc: Document object

        Yields:
            Either Paragraph or Table objects
        """
        # Based on python-docx documentation example for mixed content
        body = doc._body._body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def _extract_table(self, table: Table) -> Dict[str, Any]:
        """
        Extract data from a table

        Args:
            table: Table object

        Returns:
            Dictionary with table information
        """
        table_data = {
            "rows": [],
            "num_rows": len(table.rows),
            "num_cols": len(table.columns),
        }

        # Process table row by row
        for i, row in enumerate(table.rows):
            row_data: List[Dict[str, Any]] = []

            for j, cell in enumerate(row.cells):
                # Get cell text (joining paragraphs)
                cell_text = "\n".join(
                    [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                )

                # Get cell properties
                cell_data = {"text": cell_text, "row": i, "col": j}

                # Try to get cell formatting
                try:
                    # Add cell background color if available
                    if cell._tc.tcPr is not None and cell._tc.tcPr.shd is not None:
                        bg_color = cell._tc.tcPr.shd.val
                        if bg_color:
                            cell_data["background"] = bg_color
                except Exception as e:
                    logger.debug(
                        f"Could not extract cell formatting for cell ({i},{j}): {e}"
                    )
                    pass

                # Explicit check to help mypy
                if isinstance(row_data, list):
                    row_data.append(cell_data) # Add check and ignore potentially incorrect attr-defined

            table_data["rows"].append(row_data)

        return table_data

    def _extract_images(
        self, doc: DocxDocument, docx_data: bytes
    ) -> List[Dict[str, Any]]:
        """
        Extract images from the document

        Args:
            doc: Document object
            docx_data: Raw DOCX file data

        Returns:
            List of image information dictionaries
        """
        images = []

        try:
            # Create a temporary file to work with
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(docx_data)
                temp_docx_path = tmp.name

            # Extract images using lxml to access the docx zip contents
            # F841: Removed unused variable word_namespace
            # word_namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            #                   'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            #                   'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}
            # Define necessary namespaces
            # F841: Removed unused variable namespaces
            # namespaces = {
            #     'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            #     'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            # }

            # Find all image references in the document
            # F841: Removed unused variable rels_path
            # rels_path = 'word/_rels/document.xml.rels'
            for i, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.reltype:
                    # Get image data
                    image_part = rel.target_part
                    image_bytes = image_part.blob

                    try:
                        # Try to open the image to get dimensions
                        img = Image.open(io.BytesIO(image_bytes))
                        width, height = img.size
                        format_name = img.format
                    except Exception as e:
                        logger.warning(
                            f"Could not determine image dimensions/format for image {i + 1}: {e}"
                        )
                        width, height = 0, 0
                        format_name = "UNKNOWN"

                    # Create base64 encoded image for easy use
                    image_b64 = base64.b64encode(image_bytes).decode("utf-8")

                    # Add image info
                    image_info = {
                        "id": f"image_{i + 1}",
                        "filename": rel.target_part.partname.split("/")[-1],
                        "width": width,
                        "height": height,
                        "format": format_name,
                        "base64": image_b64[:100] + "..."
                        if len(image_b64) > 100
                        else image_b64,  # Truncate for readability
                    }

                    images.append(image_info)

            # Clean up
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)

        except Exception as e:
            logger.error(f"Error extracting images: {str(e)}")

        return images

    def _identify_document_structure(
        self, doc: DocxDocument, content: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Identify structure elements like sections, titles, etc.

        Args:
            doc: Document object
            content: Extracted content

        Returns:
            Dictionary with document structure information
        """
        structure: Dict[str, Any] = {"title": "", "sections": []}

        # Find document title
        if content["paragraphs_with_style"] and content["paragraphs_with_style"][0][
            "style"
        ] in ["Title", "Heading 1"]:
            structure["title"] = content["paragraphs_with_style"][0]["text"]

        # Build a list of all headings with their content
        current_section = None
        current_section_content: List[str] = []
        current_level = 0

        for para in content["paragraphs_with_style"]:
            if para["is_heading"]:
                # If we already have a section, save it before starting a new one
                if current_section:
                    section_text = "\n".join(current_section_content)
                    assert current_section is not None # Ensure current_section is not None before use
                    section = DocxSection(
                        heading=current_section,
                        level=current_level,
                        content=section_text,
                        paragraphs=current_section_content,
                        tables=[],  # Will fill tables later
                    )
                    structure["sections"].append(vars(section))

                # Start new section
                current_section = para["text"]
                current_level = para["heading_level"]
                current_section_content = []
            else:
                # Add paragraph to current section
                current_section_content.append(para["text"])

        # Don't forget to add the last section
        if current_section:
            section_text = "\n".join(current_section_content)
            assert current_section is not None # Ensure current_section is not None before use
            section = DocxSection(
                heading=current_section,
                level=current_level,
                content=section_text,
                paragraphs=current_section_content,
                tables=content["tables"],
            )
            structure["sections"].append(vars(section))

        # If no sections were created but we have content, create a default section
        if not structure["sections"] and content["paragraphs"]:
            default_title = structure["title"] if structure["title"] else "Document"
            section = DocxSection(
                heading=default_title,
                level=0,
                content="\n".join(content["paragraphs"]),
                paragraphs=content["paragraphs"],
                tables=content["tables"],
            )
            structure["sections"].append(vars(section))

        return structure

    def _assemble_full_text(self, paragraphs: List[str]) -> str:
        """
        Assemble full document text from paragraphs

        Args:
            paragraphs: List of paragraph texts

        Returns:
            Full document text
        """
        return "\n\n".join(paragraphs)


# Helper functions for easier usage
def parse_docx_file(file_path: str, extract_images: bool = True) -> Dict[str, Any]:
    """
    Parse DOCX file and extract information

    Args:
        file_path: Path to the DOCX file
        extract_images: Whether to extract images

    Returns:
        Dictionary with parsed information
    """
    parser = DocxParser(extract_images=extract_images)
    return parser.parse_file(file_path)


def extract_docx_text(file_path: str) -> str:
    """
    Extract only text content from a DOCX file

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    result = parse_docx_file(file_path, extract_images=False)
    if result.get("success", False):
        return result["full_text"]
    return ""


def extract_docx_tables(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract only tables from a DOCX file

    Args:
        file_path: Path to the DOCX file

    Returns:
        List of extracted tables
    """
    result = parse_docx_file(file_path, extract_images=False)
    if result.get("success", False):
        return result["content"]["tables"]
    return []
