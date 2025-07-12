"""
Document Processing Component for SutazAI Streamlit Interface
Handles PDF, DOCX, TXT, and other document formats with AI analysis
"""

import streamlit as st
import tempfile
import os
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import logging
from datetime import datetime
import asyncio
import json

# Document processing libraries
try:
    import fitz  # PyMuPDF for PDF processing
    import docx  # python-docx for DOCX processing
    import pandas as pd
    from PIL import Image
    import pytesseract  # OCR for image-based PDFs
    import textract  # For various document formats
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced Document Processing System with AI Analysis"""
    
    def __init__(self, model_manager=None, vector_memory=None):
        self.model_manager = model_manager
        self.vector_memory = vector_memory
        self.supported_formats = {
            'pdf': 'PDF Documents',
            'docx': 'Microsoft Word Documents', 
            'txt': 'Plain Text Files',
            'md': 'Markdown Files',
            'csv': 'CSV Files',
            'xlsx': 'Excel Files',
            'pptx': 'PowerPoint Files',
            'rtf': 'Rich Text Format',
            'odt': 'OpenDocument Text',
            'html': 'HTML Files'
        }
        self.processing_history = []
    
    def render_document_upload(self):
        """Render document upload interface"""
        st.header("📄 Document Upload & Processing")
        
        # File upload
        uploaded_files = st.file_uploader(
            "Upload documents for AI analysis",
            type=list(self.supported_formats.keys()),
            accept_multiple_files=True,
            help="Supported formats: " + ", ".join(self.supported_formats.values())
        )
        
        if uploaded_files:
            self._display_uploaded_files(uploaded_files)
            
            # Processing options
            st.subheader("🔧 Processing Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                extract_text = st.checkbox("Extract Text Content", value=True)
                extract_metadata = st.checkbox("Extract Metadata", value=True)
                perform_ocr = st.checkbox("Perform OCR on Images", value=False)
                
            with col2:
                generate_summary = st.checkbox("Generate AI Summary", value=True)
                extract_entities = st.checkbox("Extract Named Entities", value=True)
                sentiment_analysis = st.checkbox("Sentiment Analysis", value=False)
            
            # Advanced options
            with st.expander("🔬 Advanced Processing Options"):
                chunk_size = st.slider("Text Chunk Size", 100, 2000, 500)
                overlap_size = st.slider("Chunk Overlap", 0, 200, 50)
                language = st.selectbox("Document Language", 
                                      ["auto", "en", "es", "fr", "de", "it", "pt", "ru", "zh", "ja"])
                
                processing_mode = st.selectbox("Processing Mode",
                                             ["Standard", "Deep Analysis", "Quick Scan"])
            
            # Process button
            if st.button("🚀 Process Documents", type="primary"):
                self._process_documents(
                    uploaded_files,
                    {
                        'extract_text': extract_text,
                        'extract_metadata': extract_metadata,
                        'perform_ocr': perform_ocr,
                        'generate_summary': generate_summary,
                        'extract_entities': extract_entities,
                        'sentiment_analysis': sentiment_analysis,
                        'chunk_size': chunk_size,
                        'overlap_size': overlap_size,
                        'language': language,
                        'processing_mode': processing_mode
                    }
                )
    
    def _display_uploaded_files(self, uploaded_files: List):
        """Display information about uploaded files"""
        st.subheader("📋 Uploaded Files")
        
        for file in uploaded_files:
            with st.expander(f"📄 {file.name}", expanded=False):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.write(f"**Size:** {file.size:,} bytes")
                
                with col2:
                    st.write(f"**Type:** {file.type}")
                
                with col3:
                    st.write(f"**Format:** {Path(file.name).suffix.lower()}")
                
                # File preview
                if file.type.startswith('text/') or file.name.endswith('.txt'):
                    content_preview = str(file.read(500), 'utf-8', errors='ignore')
                    file.seek(0)  # Reset file pointer
                    st.text_area("Preview", content_preview, height=100, disabled=True)
    
    def _process_documents(self, uploaded_files: List, options: Dict):
        """Process uploaded documents with specified options"""
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        results = []
        
        for idx, file in enumerate(uploaded_files):
            status_text.text(f"Processing {file.name}...")
            progress_bar.progress((idx + 1) / len(uploaded_files))
            
            try:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.name).suffix) as tmp_file:
                    tmp_file.write(file.getvalue())
                    tmp_file_path = tmp_file.name
                
                # Process the document
                result = self._process_single_document(tmp_file_path, file.name, options)
                results.append(result)
                
                # Clean up temp file
                os.unlink(tmp_file_path)
                
            except Exception as e:
                st.error(f"Error processing {file.name}: {str(e)}")
                logger.error(f"Document processing error for {file.name}: {e}")
        
        status_text.text("Processing complete!")
        
        # Store results in session state
        if 'document_results' not in st.session_state:
            st.session_state.document_results = []
        
        st.session_state.document_results.extend(results)
        
        # Display results
        self._display_processing_results(results)
    
    def _process_single_document(self, file_path: str, file_name: str, options: Dict) -> Dict:
        """Process a single document and return results"""
        file_extension = Path(file_path).suffix.lower()
        
        result = {
            'file_name': file_name,
            'file_path': file_path,
            'processed_at': datetime.now().isoformat(),
            'file_type': file_extension,
            'content': {},
            'metadata': {},
            'analysis': {}
        }
        
        try:
            # Extract text content
            if options.get('extract_text', True):
                text_content = self._extract_text(file_path, file_extension, options)
                result['content']['text'] = text_content
                result['content']['word_count'] = len(text_content.split()) if text_content else 0
                result['content']['char_count'] = len(text_content) if text_content else 0
            
            # Extract metadata
            if options.get('extract_metadata', True):
                metadata = self._extract_metadata(file_path, file_extension)
                result['metadata'] = metadata
            
            # AI Analysis
            if options.get('generate_summary', True) and result['content'].get('text'):
                summary = self._generate_ai_summary(result['content']['text'])
                result['analysis']['summary'] = summary
            
            if options.get('extract_entities', True) and result['content'].get('text'):
                entities = self._extract_named_entities(result['content']['text'])
                result['analysis']['entities'] = entities
            
            if options.get('sentiment_analysis', False) and result['content'].get('text'):
                sentiment = self._analyze_sentiment(result['content']['text'])
                result['analysis']['sentiment'] = sentiment
            
            # Store in vector memory if available
            if self.vector_memory and result['content'].get('text'):
                self._store_in_vector_memory(result)
            
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"Error in document processing: {e}")
        
        return result
    
    def _extract_text(self, file_path: str, file_extension: str, options: Dict) -> str:
        """Extract text content from various document formats"""
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path, options)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.md']:
                return self._extract_plain_text(file_path)
            elif file_extension == '.csv':
                return self._extract_csv_text(file_path)
            elif file_extension == '.xlsx':
                return self._extract_excel_text(file_path)
            else:
                # Use textract for other formats
                return textract.process(file_path).decode('utf-8')
                
        except Exception as e:
            logger.error(f"Text extraction error for {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str, options: Dict) -> str:
        """Extract text from PDF files"""
        text_content = ""
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                # If no text found and OCR is enabled, try OCR
                if not text.strip() and options.get('perform_ocr', False):
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("ppm")
                    
                    # Save temp image for OCR
                    with tempfile.NamedTemporaryFile(suffix='.ppm', delete=False) as tmp_img:
                        tmp_img.write(img_data)
                        text = pytesseract.image_to_string(tmp_img.name)
                        os.unlink(tmp_img.name)
                
                text_content += text + "\n"
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
        
        return text_content
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            return ""
    
    def _extract_plain_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Plain text extraction error: {e}")
            return ""
    
    def _extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"CSV text extraction error: {e}")
            return ""
    
    def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel files"""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Excel text extraction error: {e}")
            return ""
    
    def _extract_metadata(self, file_path: str, file_extension: str) -> Dict:
        """Extract metadata from documents"""
        metadata = {
            'file_size': os.path.getsize(file_path),
            'created_at': datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            'modified_at': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
        }
        
        try:
            if file_extension == '.pdf':
                doc = fitz.open(file_path)
                pdf_metadata = doc.metadata
                metadata.update({
                    'title': pdf_metadata.get('title', ''),
                    'author': pdf_metadata.get('author', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'creation_date': pdf_metadata.get('creationDate', ''),
                    'modification_date': pdf_metadata.get('modDate', ''),
                    'page_count': doc.page_count
                })
                doc.close()
                
            elif file_extension == '.docx':
                doc = docx.Document(file_path)
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                })
                
        except Exception as e:
            logger.error(f"Metadata extraction error: {e}")
        
        return metadata
    
    def _generate_ai_summary(self, text: str) -> str:
        """Generate AI summary of the document"""
        try:
            if self.model_manager:
                # Use the model manager to generate summary
                prompt = f"Please provide a concise summary of the following text:\n\n{text[:2000]}..."
                summary = self.model_manager.generate_response(prompt)
                return summary
            else:
                # Fallback: simple extractive summary
                sentences = text.split('.')[:5]
                return '. '.join(sentences) + '.'
                
        except Exception as e:
            logger.error(f"AI summary generation error: {e}")
            return "Summary generation failed."
    
    def _extract_named_entities(self, text: str) -> List[Dict]:
        """Extract named entities from text"""
        try:
            # This would typically use spaCy or similar NLP library
            # For now, return mock entities
            entities = [
                {"text": "Sample Entity", "label": "PERSON", "start": 0, "end": 13},
                {"text": "Organization", "label": "ORG", "start": 20, "end": 32}
            ]
            return entities
            
        except Exception as e:
            logger.error(f"Named entity extraction error: {e}")
            return []
    
    def _analyze_sentiment(self, text: str) -> Dict:
        """Analyze sentiment of the document"""
        try:
            # Mock sentiment analysis - would use actual NLP library
            import random
            
            sentiment_score = random.uniform(-1, 1)
            
            if sentiment_score > 0.1:
                label = "positive"
            elif sentiment_score < -0.1:
                label = "negative"
            else:
                label = "neutral"
            
            return {
                "label": label,
                "score": sentiment_score,
                "confidence": random.uniform(0.7, 0.95)
            }
            
        except Exception as e:
            logger.error(f"Sentiment analysis error: {e}")
            return {"label": "unknown", "score": 0, "confidence": 0}
    
    def _store_in_vector_memory(self, result: Dict):
        """Store document content in vector memory"""
        try:
            if self.vector_memory and result['content'].get('text'):
                # Create chunks and store in vector memory
                text = result['content']['text']
                chunks = self._chunk_text(text, 500, 50)
                
                for i, chunk in enumerate(chunks):
                    self.vector_memory.store_memory(
                        f"{result['file_name']}_chunk_{i}",
                        chunk,
                        metadata={
                            'source_file': result['file_name'],
                            'chunk_index': i,
                            'total_chunks': len(chunks),
                            'file_type': result['file_type']
                        }
                    )
                    
        except Exception as e:
            logger.error(f"Vector memory storage error: {e}")
    
    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            
            if i + chunk_size >= len(words):
                break
        
        return chunks
    
    def _display_processing_results(self, results: List[Dict]):
        """Display processing results"""
        st.success(f"✅ Successfully processed {len(results)} documents!")
        
        for result in results:
            with st.expander(f"📄 {result['file_name']} Results", expanded=True):
                
                # Basic info
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Word Count", result['content'].get('word_count', 0))
                
                with col2:
                    st.metric("Character Count", result['content'].get('char_count', 0))
                
                with col3:
                    st.metric("File Size", f"{result['metadata'].get('file_size', 0):,} bytes")
                
                # Content tabs
                tab1, tab2, tab3, tab4 = st.tabs(["📝 Content", "🔍 Analysis", "📊 Metadata", "💾 Export"])
                
                with tab1:
                    if result['content'].get('text'):
                        st.text_area(
                            "Extracted Text",
                            result['content']['text'][:1000] + "..." if len(result['content']['text']) > 1000 else result['content']['text'],
                            height=300,
                            disabled=True
                        )
                    else:
                        st.info("No text content extracted")
                
                with tab2:
                    if result['analysis'].get('summary'):
                        st.subheader("📋 AI Summary")
                        st.write(result['analysis']['summary'])
                    
                    if result['analysis'].get('entities'):
                        st.subheader("🏷️ Named Entities")
                        for entity in result['analysis']['entities']:
                            st.write(f"• **{entity['text']}** ({entity['label']})")
                    
                    if result['analysis'].get('sentiment'):
                        st.subheader("😊 Sentiment Analysis")
                        sentiment = result['analysis']['sentiment']
                        st.write(f"**Sentiment:** {sentiment['label'].title()}")
                        st.write(f"**Score:** {sentiment['score']:.2f}")
                        st.write(f"**Confidence:** {sentiment['confidence']:.2%}")
                
                with tab3:
                    st.subheader("📊 Document Metadata")
                    metadata_df = pd.DataFrame(list(result['metadata'].items()), columns=['Property', 'Value'])
                    st.dataframe(metadata_df, use_container_width=True)
                
                with tab4:
                    st.subheader("💾 Export Options")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if st.button(f"📄 Export Text", key=f"export_text_{result['file_name']}"):
                            self._export_text(result)
                    
                    with col2:
                        if st.button(f"📊 Export Analysis", key=f"export_analysis_{result['file_name']}"):
                            self._export_analysis(result)
                    
                    with col3:
                        if st.button(f"🗂️ Export All", key=f"export_all_{result['file_name']}"):
                            self._export_complete_result(result)
    
    def _export_text(self, result: Dict):
        """Export extracted text"""
        if result['content'].get('text'):
            st.download_button(
                label="Download Text Content",
                data=result['content']['text'],
                file_name=f"{Path(result['file_name']).stem}_extracted.txt",
                mime="text/plain"
            )
    
    def _export_analysis(self, result: Dict):
        """Export analysis results"""
        analysis_data = json.dumps(result['analysis'], indent=2)
        st.download_button(
            label="Download Analysis Results",
            data=analysis_data,
            file_name=f"{Path(result['file_name']).stem}_analysis.json",
            mime="application/json"
        )
    
    def _export_complete_result(self, result: Dict):
        """Export complete processing result"""
        complete_data = json.dumps(result, indent=2, default=str)
        st.download_button(
            label="Download Complete Results",
            data=complete_data,
            file_name=f"{Path(result['file_name']).stem}_complete.json",
            mime="application/json"
        )
    
    def render_processing_history(self):
        """Render document processing history"""
        st.header("📚 Processing History")
        
        if 'document_results' not in st.session_state or not st.session_state.document_results:
            st.info("No documents have been processed yet.")
            return
        
        results = st.session_state.document_results
        
        # Summary statistics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Documents", len(results))
        
        with col2:
            total_words = sum(r['content'].get('word_count', 0) for r in results)
            st.metric("Total Words", f"{total_words:,}")
        
        with col3:
            total_chars = sum(r['content'].get('char_count', 0) for r in results)
            st.metric("Total Characters", f"{total_chars:,}")
        
        with col4:
            file_types = len(set(r['file_type'] for r in results))
            st.metric("File Types", file_types)
        
        # Results table
        st.subheader("📋 Document List")
        
        history_data = []
        for result in results:
            history_data.append({
                'File Name': result['file_name'],
                'Type': result['file_type'],
                'Words': result['content'].get('word_count', 0),
                'Processed': result['processed_at'][:19],  # Remove microseconds
                'Has Summary': '✅' if result['analysis'].get('summary') else '❌',
                'Has Entities': '✅' if result['analysis'].get('entities') else '❌'
            })
        
        df = pd.DataFrame(history_data)
        st.dataframe(df, use_container_width=True)
        
        # Clear history button
        if st.button("🗑️ Clear Processing History"):
            st.session_state.document_results = []
            st.success("Processing history cleared!")
            st.rerun()
    
    def render_batch_processing(self):
        """Render batch processing interface"""
        st.header("⚡ Batch Processing")
        st.info("Upload multiple documents for automated batch processing with predefined settings.")
        
        # Predefined processing profiles
        profiles = {
            "Standard Analysis": {
                "extract_text": True,
                "extract_metadata": True,
                "generate_summary": True,
                "extract_entities": True,
                "sentiment_analysis": False,
                "perform_ocr": False
            },
            "Deep Research": {
                "extract_text": True,
                "extract_metadata": True,
                "generate_summary": True,
                "extract_entities": True,
                "sentiment_analysis": True,
                "perform_ocr": True
            },
            "Quick Scan": {
                "extract_text": True,
                "extract_metadata": False,
                "generate_summary": False,
                "extract_entities": False,
                "sentiment_analysis": False,
                "perform_ocr": False
            }
        }
        
        selected_profile = st.selectbox("Select Processing Profile", list(profiles.keys()))
        
        # Display profile settings
        with st.expander("📋 Profile Settings"):
            for setting, value in profiles[selected_profile].items():
                st.write(f"• **{setting.replace('_', ' ').title()}**: {'✅' if value else '❌'}")
        
        # Batch upload
        batch_files = st.file_uploader(
            "Upload documents for batch processing",
            type=list(self.supported_formats.keys()),
            accept_multiple_files=True,
            key="batch_upload"
        )
        
        if batch_files:
            st.write(f"📁 **{len(batch_files)} files ready for batch processing**")
            
            if st.button("🚀 Start Batch Processing", type="primary"):
                batch_options = profiles[selected_profile].copy()
                batch_options.update({
                    'chunk_size': 500,
                    'overlap_size': 50,
                    'language': 'auto',
                    'processing_mode': 'Standard'
                })
                
                self._process_documents(batch_files, batch_options)
    
    def render(self):
        """Render the complete document processor"""
        tab1, tab2, tab3, tab4 = st.tabs([
            "📤 Upload & Process",
            "📚 Processing History", 
            "⚡ Batch Processing",
            "ℹ️ Supported Formats"
        ])
        
        with tab1:
            self.render_document_upload()
        
        with tab2:
            self.render_processing_history()
        
        with tab3:
            self.render_batch_processing()
        
        with tab4:
            st.header("📋 Supported Document Formats")
            
            format_data = []
            for ext, desc in self.supported_formats.items():
                format_data.append({
                    'Extension': f'.{ext}',
                    'Description': desc,
                    'Text Extraction': '✅',
                    'Metadata Extraction': '✅' if ext in ['pdf', 'docx'] else '⚠️',
                    'OCR Support': '✅' if ext == 'pdf' else '❌'
                })
            
            df = pd.DataFrame(format_data)
            st.dataframe(df, use_container_width=True)
            
            st.info("💡 **Tip**: For best results with image-based PDFs, enable OCR processing.")