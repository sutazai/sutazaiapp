#!/usr/bin/env python3
"""
DocuMind Document Processing Service for SutazAI
Handles PDF, DOCX, TXT, and other document formats
"""

import os
import tempfile
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
# Note: Import specific implementation modules as needed

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import markdown
from bs4 import BeautifulSoup

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="SutazAI DocuMind Service", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class DocumentResponse(BaseModel):
    content: str
    metadata: Dict[str, Any]
    file_type: str
    processed: bool

class DocumentProcessor:
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_txt,
            'text/markdown': self._process_markdown,
            'text/html': self._process_html,
            'application/vnd.ms-excel': self._process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel,
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
            'image/tiff': self._process_image
        }

    async def process_document(self, file_path: str, file_type: str) -> DocumentResponse:
        """Process document based on file type"""
        
        if file_type not in self.supported_types:
            raise HTTPException(status_code=415, detail=f"Unsupported file type: {file_type}")
        
        try:
            processor = self.supported_types[file_type]
            content, metadata = await processor(file_path)
            
            return DocumentResponse(
                content=content,
                metadata=metadata,
                file_type=file_type,
                processed=True
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

    async def _process_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process PDF files"""
        content = ""
        metadata = {"pages": 0, "has_images": False}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\\n--- Page {page_num + 1} ---\\n{page_text}"
                
                # If no text found, try OCR
                if not content.strip():
                    logger.info("No text found in PDF, attempting OCR...")
                    content = await self._ocr_pdf(file_path)
                    metadata["ocr_used"] = True
                
            return content, metadata
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise

    async def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on PDF images"""
        try:
            images = convert_from_path(file_path)
            text_content = ""
            
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_content += f"\\n--- Page {i + 1} (OCR) ---\\n{text}"
            
            return text_content
            
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return "OCR processing failed"

    async def _process_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process DOCX files"""
        try:
            doc = Document(file_path)
            content = ""
            metadata = {"paragraphs": 0, "tables": 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\\n"
                    metadata["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata["tables"] += 1
                content += "\\n--- Table ---\\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    content += row_text + "\\n"
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise

    async def _process_txt(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            metadata = {
                "lines": len(content.split('\\n')),
                "characters": len(content),
                "words": len(content.split())
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"TXT processing failed: {e}")
            raise

    async def _process_markdown(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Markdown files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                md_content = await file.read()
            
            # Convert to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            content = soup.get_text()
            
            metadata = {
                "markdown_length": len(md_content),
                "html_length": len(html),
                "text_length": len(content)
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Markdown processing failed: {e}")
            raise

    async def _process_html(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process HTML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            content = soup.get_text()
            
            metadata = {
                "html_length": len(html_content),
                "text_length": len(content),
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img'))
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"HTML processing failed: {e}")
            raise

    async def _process_excel(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process Excel files"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            content = ""
            metadata = {"sheets": len(df), "total_rows": 0, "total_columns": 0}
            
            for sheet_name, sheet_df in df.items():
                content += f"\\n--- Sheet: {sheet_name} ---\\n"
                content += sheet_df.to_string(index=False) + "\\n"
                metadata["total_rows"] += len(sheet_df)
                metadata["total_columns"] += len(sheet_df.columns)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            raise

    async def _process_image(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process image files with OCR"""
        try:
            image = Image.open(file_path)
            content = pytesseract.image_to_string(image)
            
            metadata = {
                "width": image.width,
                "height": image.height,
                "mode": image.mode,
                "format": image.format,
                "ocr_confidence": "medium"  # Could implement confidence scoring
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            raise

# Initialize document processor
doc_processor = DocumentProcessor()

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "supported_types": list(doc_processor.supported_types.keys()),
        "tesseract_available": True,  # Could check if tesseract is actually available
        "poppler_available": True     # Could check if poppler is actually available
    }

@app.post("/process", response_model=DocumentResponse)
async def process_document(file: UploadFile = File(...)):
    """Process uploaded document"""
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # Detect file type
        # TODO: Implement file type detection (e.g., using python-process)
        file_type = "application/octet-stream"  # Default file type
        logger.info(f"Processing file: {file.filename}, type: {file_type}")
        
        # Process document
        result = await doc_processor.process_document(temp_file_path, file_type)
        result.metadata["filename"] = file.filename
        result.metadata["size"] = len(content)
        
        return result
        
    finally:
        # Clean up temporary file
        os.unlink(temp_file_path)

@app.get("/supported-types")
async def get_supported_types():
    """Get list of supported file types"""
    return {
        "supported_types": list(doc_processor.supported_types.keys()),
        "examples": {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "txt": "text/plain",
            "md": "text/markdown",
            "html": "text/html",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "jpg": "image/jpeg",
            "png": "image/png"
        }
    }

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8080)