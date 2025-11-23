"""
File Upload and Document Processing Components
Handles file uploads, parsing, and context injection for RAG
"""

import streamlit as st
import io
import os
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import hashlib
import mimetypes

# Try to import document parsers
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class FileUploadHandler:
    """
    Advanced file upload handler with drag-drop,
    validation, and document parsing
    """
    
    ALLOWED_TYPES = {
        "text": [".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".xml", ".yaml", ".yml"],
        "document": [".pdf", ".docx", ".doc"],
        "data": [".csv", ".xlsx", ".xls", ".json"],
        "code": [".py", ".js", ".java", ".cpp", ".c", ".go", ".rs", ".ts"],
        "audio": [".wav", ".mp3", ".ogg", ".m4a"],
        "image": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    @staticmethod
    def render_upload_zone(accept_multiple: bool = True,
                          allowed_categories: Optional[List[str]] = None) -> List[Any]:
        """
        Render file upload zone with drag-drop support
        
        Args:
            accept_multiple: Allow multiple file uploads
            allowed_categories: List of allowed file categories (text, document, data, etc.)
        
        Returns:
            List of uploaded file objects
        """
        st.markdown("### 📁 File Upload")
        
        # Determine allowed file types
        if allowed_categories:
            allowed_extensions = []
            for category in allowed_categories:
                if category in FileUploadHandler.ALLOWED_TYPES:
                    allowed_extensions.extend(FileUploadHandler.ALLOWED_TYPES[category])
        else:
            # Allow all types
            allowed_extensions = []
            for exts in FileUploadHandler.ALLOWED_TYPES.values():
                allowed_extensions.extend(exts)
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Choose files to upload",
            type=[ext.lstrip('.') for ext in allowed_extensions],
            accept_multiple_files=accept_multiple,
            help=f"Maximum file size: {FileUploadHandler.MAX_FILE_SIZE / 1024 / 1024:.0f}MB",
            key="file_uploader_main"
        )
        
        if uploaded_files:
            # Handle single file or list
            files = uploaded_files if isinstance(uploaded_files, list) else [uploaded_files]
            
            # Display uploaded files
            st.success(f"✅ {len(files)} file(s) uploaded")
            
            # Show file details
            with st.expander("📄 Uploaded Files", expanded=True):
                for file in files:
                    FileUploadHandler._display_file_info(file)
            
            return files
        
        return []
    
    @staticmethod
    def _display_file_info(file):
        """Display information about uploaded file"""
        col1, col2, col3 = st.columns([2, 1, 1])
        
        with col1:
            st.write(f"**{file.name}**")
        
        with col2:
            file_size = file.size
            if file_size < 1024:
                size_str = f"{file_size} B"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size / 1024:.1f} KB"
            else:
                size_str = f"{file_size / 1024 / 1024:.1f} MB"
            st.write(size_str)
        
        with col3:
            file_type = file.type or "unknown"
            st.write(file_type.split('/')[-1])
    
    @staticmethod
    def parse_file(file) -> Dict[str, Any]:
        """
        Parse uploaded file and extract content
        
        Returns:
            Dictionary with parsed content and metadata
        """
        file_extension = os.path.splitext(file.name)[1].lower()
        
        result = {
            "filename": file.name,
            "extension": file_extension,
            "size": file.size,
            "type": file.type,
            "content": "",
            "error": None,
            "metadata": {}
        }
        
        try:
            # Text files
            if file_extension in FileUploadHandler.ALLOWED_TYPES["text"]:
                content = file.read().decode('utf-8')
                result["content"] = content
                result["metadata"]["lines"] = len(content.split('\n'))
                result["metadata"]["chars"] = len(content)
            
            # PDF files
            elif file_extension == ".pdf" and PDF_AVAILABLE:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                result["content"] = text
                result["metadata"]["pages"] = len(pdf_reader.pages)
            
            # DOCX files
            elif file_extension == ".docx" and DOCX_AVAILABLE:
                doc = docx.Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
                result["content"] = text
                result["metadata"]["paragraphs"] = len(doc.paragraphs)
            
            # CSV files
            elif file_extension == ".csv" and PANDAS_AVAILABLE:
                df = pd.read_csv(file)
                result["content"] = df.to_string()
                result["metadata"]["rows"] = len(df)
                result["metadata"]["columns"] = len(df.columns)
                result["metadata"]["data"] = df
            
            # JSON files
            elif file_extension == ".json":
                import json
                content = file.read().decode('utf-8')
                data = json.loads(content)
                result["content"] = json.dumps(data, indent=2)
                result["metadata"]["json_data"] = data
            
            # Code files
            elif file_extension in FileUploadHandler.ALLOWED_TYPES["code"]:
                content = file.read().decode('utf-8')
                result["content"] = content
                result["metadata"]["lines"] = len(content.split('\n'))
                result["metadata"]["language"] = file_extension.lstrip('.')
            
            else:
                result["error"] = f"Unsupported file type: {file_extension}"
        
        except Exception as e:
            result["error"] = f"Error parsing file: {str(e)}"
        
        return result
    
    @staticmethod
    def create_document_context(parsed_files: List[Dict[str, Any]]) -> str:
        """
        Create combined context from multiple documents for RAG
        
        Args:
            parsed_files: List of parsed file dictionaries
        
        Returns:
            Combined document context string
        """
        context_parts = []
        
        context_parts.append("=== DOCUMENT CONTEXT ===\n")
        context_parts.append(f"Total Documents: {len(parsed_files)}\n")
        context_parts.append(f"Uploaded: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        context_parts.append("\n")
        
        for i, doc in enumerate(parsed_files, 1):
            context_parts.append(f"--- Document {i}: {doc['filename']} ---\n")
            
            if doc.get("error"):
                context_parts.append(f"ERROR: {doc['error']}\n")
            else:
                # Add metadata
                if doc.get("metadata"):
                    meta = doc["metadata"]
                    context_parts.append(f"Metadata: {', '.join([f'{k}={v}' for k, v in meta.items() if k != 'data' and k != 'json_data'])}\n")
                
                # Add content (truncate if too long)
                content = doc.get("content", "")
                if len(content) > 5000:
                    context_parts.append(f"{content[:5000]}...\n[Content truncated]\n")
                else:
                    context_parts.append(f"{content}\n")
            
            context_parts.append("\n")
        
        context_parts.append("=== END DOCUMENT CONTEXT ===\n")
        
        return "".join(context_parts)
    
    @staticmethod
    def render_document_preview(parsed_file: Dict[str, Any]):
        """Render preview of parsed document"""
        st.markdown(f"### 📄 {parsed_file['filename']}")
        
        if parsed_file.get("error"):
            st.error(f"❌ {parsed_file['error']}")
            return
        
        # Metadata
        if parsed_file.get("metadata"):
            cols = st.columns(len(parsed_file["metadata"]))
            for i, (key, value) in enumerate(parsed_file["metadata"].items()):
                if key not in ["data", "json_data"]:
                    with cols[i]:
                        st.metric(key.title(), value)
        
        # Content preview
        content = parsed_file.get("content", "")
        
        # Special handling for DataFrames
        if "data" in parsed_file.get("metadata", {}):
            st.dataframe(parsed_file["metadata"]["data"], use_container_width=True)
        
        # Code files - syntax highlighting
        elif parsed_file["extension"] in FileUploadHandler.ALLOWED_TYPES["code"]:
            language = parsed_file.get("metadata", {}).get("language", "python")
            st.code(content, language=language)
        
        # Regular text/documents
        else:
            # Truncate long content
            if len(content) > 1000:
                st.text_area("Content Preview", content[:1000] + "\n...[truncated]", height=300)
                st.caption(f"Showing first 1000 of {len(content)} characters")
            else:
                st.text_area("Content Preview", content, height=300)


class DocumentProcessor:
    """Process documents for RAG and context injection"""
    
    @staticmethod
    def chunk_document(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split document into overlapping chunks for better context
        
        Args:
            text: Document text
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks
        
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to find a sentence boundary
            if end < len(text):
                # Look for period, question mark, or exclamation
                boundary = max(
                    text.rfind('.', start, end),
                    text.rfind('?', start, end),
                    text.rfind('!', start, end)
                )
                
                if boundary > start:
                    end = boundary + 1
            
            chunks.append(text[start:end].strip())
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
        """
        Extract keywords from text using simple frequency analysis
        
        Args:
            text: Document text
            max_keywords: Maximum number of keywords to extract
        
        Returns:
            List of keywords
        """
        # Simple word frequency (can be enhanced with NLP)
        words = text.lower().split()
        
        # Filter out common stopwords
        stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                    'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
                    'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they'}
        
        # Count word frequencies
        word_freq = {}
        for word in words:
            # Basic cleaning
            word = ''.join(c for c in word if c.isalnum())
            if len(word) > 3 and word not in stopwords:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        return [word for word, freq in sorted_words[:max_keywords]]
    
    @staticmethod
    def calculate_document_hash(content: str) -> str:
        """Calculate unique hash for document content"""
        return hashlib.sha256(content.encode()).hexdigest()[:16]
    
    @staticmethod
    def format_for_rag(parsed_files: List[Dict[str, Any]], 
                      include_metadata: bool = True) -> List[Dict[str, Any]]:
        """
        Format documents for RAG (Retrieval-Augmented Generation)
        
        Returns:
            List of formatted document chunks ready for vector database
        """
        formatted_docs = []
        
        for doc in parsed_files:
            if doc.get("error"):
                continue
            
            content = doc.get("content", "")
            
            # Chunk the document
            chunks = DocumentProcessor.chunk_document(content)
            
            # Extract keywords
            keywords = DocumentProcessor.extract_keywords(content)
            
            # Create document entries
            for i, chunk in enumerate(chunks):
                entry = {
                    "text": chunk,
                    "chunk_id": i,
                    "total_chunks": len(chunks),
                    "source": doc["filename"],
                    "doc_hash": DocumentProcessor.calculate_document_hash(content)
                }
                
                if include_metadata and doc.get("metadata"):
                    entry["metadata"] = doc["metadata"]
                
                if i == 0:
                    # Add keywords only to first chunk
                    entry["keywords"] = keywords
                
                formatted_docs.append(entry)
        
        return formatted_docs
